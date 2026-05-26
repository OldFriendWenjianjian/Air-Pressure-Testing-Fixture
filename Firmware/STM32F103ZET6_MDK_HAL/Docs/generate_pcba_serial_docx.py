from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "PCBA串口对接说明.docx"


def crc16_modbus(data: bytes) -> int:
    crc = 0xFFFF
    for b in data:
        crc ^= b
        for _ in range(8):
            if crc & 1:
                crc = (crc >> 1) ^ 0xA001
            else:
                crc >>= 1
    return crc & 0xFFFF


def frame(cmd: int, ch: int, data: bytes = b"") -> str:
    body = bytes([cmd, ch, len(data) & 0xFF, (len(data) >> 8) & 0xFF]) + data
    crc = crc16_modbus(body)
    raw = bytes([0x55, 0xAA]) + body + bytes([crc & 0xFF, (crc >> 8) & 0xFF])
    return " ".join(f"{b:02X}" for b in raw)


def unknown_frame(cmd: int, ch: int, data_len: int) -> str:
    prefix = f"55 AA {cmd:02X} {ch:02X} {data_len & 0xFF:02X} {(data_len >> 8) & 0xFF:02X}"
    data = " ".join("??" for _ in range(data_len))
    return f"{prefix} {data} ?? ??"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, mono: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8 if mono else 9)
    if mono:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def style_table(table, header_fill: str = "D9EAF7") -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    style_table(table)


def add_command_table(doc: Document) -> None:
    rows = [
        ("00", "唤醒机器", "MCU -> PCBA", "无", "空应答包 7F"),
        ("01", "进入测试模式", "MCU -> PCBA", "无", "空应答包 7F"),
        ("02", "记录零点", "MCU -> PCBA", "无", "空应答包 7F"),
        ("03", "开机", "MCU -> PCBA", "无", "空应答包 7F"),
        ("04", "查询低电状态", "MCU -> PCBA", "无", "7F + 1字节，低电=00，非低电=01"),
        ("05", "查询正常状态", "MCU -> PCBA", "无", "7F + 1字节，低电=01，非低电=00"),
        ("10", "标定压力同步", "MCU -> PCBA", "4字节压力值", "空应答包 7F"),
        ("11", "测试压力查询", "MCU -> PCBA", "无", "11 + 4字节 PCBA 测量值"),
        ("7F", "应答包", "PCBA -> MCU", "无或1字节", "PCBA 对主控命令的应答"),
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["命令", "名称", "方向", "DATA", "说明/返回"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)


def add_stage_table(doc: Document) -> None:
    rows = [
        ("U盘维护模式", "无 PCBA 业务命令", "开机长按 KEY1(PC3)，USB MSC 暴露 W25Q128"),
        ("初始化模式，罐体压力闭环", "无 PCBA 业务命令", "主控只进行 6 个气罐压力库存闭环"),
        ("自动气密性测试", "无 PCBA 业务命令", "主控关闭阀门并保压监测"),
        ("Ready 等待压合", "无 PCBA 业务命令", "等待外置压合开关"),
        ("Ready-压合后第1步", "00 唤醒机器", "单通道等待 2s，PCBA 回空应答包"),
        ("Ready-压合后第2步", "01 进入测试模式", "常规等待 10ms，PCBA 回空应答包"),
        ("Ready-压合后第3步", "03 开机", "常规等待 10ms，PCBA 回空应答包，随后主控等待 2s"),
        ("Ready-压合后第4步", "02 记录零点", "常规等待 10ms，PCBA 回空应答包"),
        ("Ready-压合后第5步", "04 查询低电状态", "先切到 4.5V 并等待 0.5s，期望返回 00"),
        ("Ready-压合后第6步", "05 查询正常状态", "先切回 5V 并等待 0.5s，期望返回 00"),
        ("50mmHg 标定", "10 + 当前真实压力值", "压力值未知时写 ?? ?? ?? ??，CRC 同步为 ?? ??"),
        ("150mmHg 标定", "10 + 当前真实压力值", "压力值未知时写 ?? ?? ?? ??，CRC 同步为 ?? ??"),
        ("250mmHg 标定", "10 + 当前真实压力值", "压力值未知时写 ?? ?? ?? ??，CRC 同步为 ?? ??"),
        ("100mmHg 测试", "11 测试压力查询", "PCBA 返回 4 字节当前压力测量值"),
        ("200mmHg 测试", "11 测试压力查询", "PCBA 返回 4 字节当前压力测量值"),
        ("300mmHg 测试", "11 测试压力查询", "PCBA 返回 4 字节当前压力测量值"),
        ("结果显示/补气", "无 PCBA 业务命令", "主控汇总结果并准备下一批次"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["阶段", "MCU 输出", "时序/PCBA 响应"]):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)


def add_fixed_frame_table(doc: Document, commands: list[tuple[str, int]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    headers = ["阶段/命令", "通道", "MCU 输出完整帧"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for name, cmd in commands:
        for ch in range(1, 9):
            cells = table.add_row().cells
            cells[0].text = name
            cells[1].text = f"CH{ch}"
            set_cell_text(cells[2], frame(cmd, ch), mono=True)
    style_table(table)


def add_unknown_pressure_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["阶段", "通道", "MCU 输出帧（压力未知时）"]):
        table.rows[0].cells[i].text = h
    for stage in ["50mmHg 标定", "150mmHg 标定", "250mmHg 标定"]:
        for ch in range(1, 9):
            cells = table.add_row().cells
            cells[0].text = stage
            cells[1].text = f"CH{ch}"
            set_cell_text(cells[2], unknown_frame(0x10, ch, 4), mono=True)
    style_table(table)


def add_pcba_response_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["响应类型", "PCBA 输出帧", "CH1 样例", "说明"]):
        table.rows[0].cells[i].text = h
    response_rows = [
        ("空应答包", "55 AA 7F CH 00 00 CRC_L CRC_H", frame(0x7F, 1), "用于 00/01/02/03/10 成功应答；CH 换成实际通道后 CRC 需重算"),
        ("一字节 YES/期望值", "55 AA 7F CH 01 00 00 CRC_L CRC_H", frame(0x7F, 1, bytes([0x00])), "低电查询期望低电=00；正常查询期望非低电=00"),
        ("一字节 NO/非期望值", "55 AA 7F CH 01 00 01 CRC_L CRC_H", frame(0x7F, 1, bytes([0x01])), "低电查询非低电=01；正常查询低电=01"),
        ("测试压力返回", "55 AA 11 CH 04 00 ?? ?? ?? ?? ?? ??", unknown_frame(0x11, 1, 4), "DATA 为 PCBA 当前压力测量值，未知量用 ??；填写真实测量值后 CRC 需重算"),
    ]
    for row in response_rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        set_cell_text(cells[1], row[1], mono=True)
        set_cell_text(cells[2], row[2], mono=True)
        cells[3].text = row[3]
    style_table(table)


def set_document_font(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build() -> None:
    doc = Document()
    set_document_font(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("气压检测工装 PCBA 串口对接说明")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("依据：气压检测工装功能逻辑表.md 第 4 节与第 4.1 节").font.size = Pt(10)

    add_heading(doc, "1. 通信规则", 1)
    add_kv_table(doc, [
        ("主从关系", "MCU 为主机，8 路 PCBA 为从机；只能由 MCU 发起，PCBA 收到指令后立即返回。"),
        ("通道映射", "UART1 至 UART8 对应 PCBA 通道 1 至通道 8。"),
        ("串口参数", "115200 8N1：115200 baud、8 数据位、无校验、1 停止位。"),
        ("轮询方式", "MCU 依次轮询 CH1 至 CH8。"),
        ("超时", "常规命令单路最大等待 10ms；压合后的第一条唤醒命令 00 单路等待 2s。"),
        ("压力单位", "气压值为 uint32 little-endian，单位 0.001mmHg；未知实时压力值在本文档中用 ?? 表示。"),
    ])

    add_heading(doc, "2. 帧格式", 1)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["字段", "字节数", "说明"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("55 AA", "2", "固定帧头"),
        ("CMD", "1", "命令字"),
        ("CH", "1", "通道号，01 至 08"),
        ("LEN_L LEN_H", "2", "DATA 长度，小端"),
        ("DATA", "N", "参数或返回数据"),
        ("CRC_L CRC_H", "2", "CRC16/MODBUS，小端"),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_para(doc, "CRC 校验范围为 CMD + CH + LEN_L + LEN_H + DATA，不包含帧头 55 AA。CRC 初值 0xFFFF，多项式 0xA001。")

    add_heading(doc, "3. 命令字", 1)
    add_command_table(doc)

    add_heading(doc, "4. 状态阶段串口交互", 1)
    add_stage_table(doc)

    add_heading(doc, "5. MCU 固定命令完整输出帧", 1)
    add_para(doc, "下表列出所有不带 DATA 的 MCU 固定命令完整帧，CRC 已按通道计算，可直接用于抓包核对。")
    add_fixed_frame_table(doc, [
        ("唤醒机器", 0x00),
        ("进入测试模式", 0x01),
        ("记录零点", 0x02),
        ("开机", 0x03),
    ])
    doc.add_page_break()
    add_heading(doc, "5. MCU 固定命令完整输出帧（续）", 1)
    add_fixed_frame_table(doc, [
        ("查询低电状态", 0x04),
        ("查询正常状态", 0x05),
        ("测试压力查询", 0x11),
    ])

    add_heading(doc, "6. MCU 标定阶段压力下发帧", 1)
    add_para(doc, "标定阶段 DATA 为主控当前读取到的真实压力值。由于该值运行时才确定，DATA 与 CRC 在本文档中用 ?? 表示。")
    add_unknown_pressure_table(doc)

    add_heading(doc, "7. PCBA 应返回的数据", 1)
    add_pcba_response_table(doc)

    add_heading(doc, "8. PCBA 侧注意事项", 1)
    for text in [
        "PCBA 仅在收到完整且 CRC 正确的帧后回包。",
        "PCBA 回包需要保持与收到命令一致的 CH。",
        "00/01/02/03/10 成功后返回空应答包。",
        "04/05 返回 7F 应答包，LEN 必须为 01 00。",
        "11 收到后返回命令字 11，LEN 为 04 00，DATA 为 PCBA 当前压力测量值。",
    ]:
        doc.add_paragraph(text, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
